# <!-- Powered by FVTeamOfAgents Core -->
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
import re

def create_docx_from_md(md_text, docx_path):
    """Enhanced Markdown to Word converter with professional formatting"""
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    lines = md_text.split('\n')
    i = 0
    in_table = False
    table_headers = []
    in_code_block = False
    code_lines = []
    
    while i < len(lines):
        line = lines[i].strip()
        
        # Handle code blocks
        if line.startswith('```'):
            if in_code_block:
                # End of code block
                if code_lines:
                    add_code_block(doc, '\n'.join(code_lines))
                    code_lines = []
                in_code_block = False
            else:
                # Start of code block
                in_code_block = True
            i += 1
            continue
            
        if in_code_block:
            code_lines.append(lines[i])
            i += 1
            continue
        
        # Handle headers
        if line.startswith('# '):
            add_styled_heading(doc, line[2:], level=1)
        elif line.startswith('## '):
            add_styled_heading(doc, line[3:], level=2)
        elif line.startswith('### '):
            add_styled_heading(doc, line[4:], level=3)
        elif line.startswith('#### '):
            add_styled_heading(doc, line[5:], level=4)
        
        # Handle tables
        elif '|' in line and line.strip().startswith('|'):
            if not in_table:
                # Start of table - extract headers
                table_headers = [cell.strip() for cell in line.split('|')[1:-1]]
                in_table = True
                table_data = []
            elif line.strip().startswith('|---') or line.strip().startswith('|-'):
                # Table separator line - skip
                pass
            else:
                # Table data row
                row_data = [cell.strip() for cell in line.split('|')[1:-1]]
                table_data.append(row_data)
        
        # Handle end of table or process non-table content
        elif in_table:
            # End of table - create the table
            if table_headers and table_data:
                add_formatted_table(doc, table_headers, table_data)
            in_table = False
            table_headers = []
            table_data = []
            
            # Process current line as regular content
            if line:
                add_formatted_paragraph(doc, line)
        
        # Handle lists and table of contents
        elif line.startswith('- ') or line.startswith('* '):
            # Check if this is a TOC sub-item (like "   - 1.1 [Introduction](#introduction)")
            if re.match(r'^\s*- \d+\.\d+', line):
                # TOC sub-item
                match = re.match(r'^\s*- (\d+\.\d+)\s+\[([^\]]+)\]', line)
                if match:
                    number = match.group(1)
                    title = match.group(2)
                    add_toc_sub_item(doc, f"{number} {title}")
                else:
                    # Fallback parsing
                    add_toc_sub_item(doc, line.strip()[2:])
            else:
                # Regular bullet point
                add_bullet_point(doc, line[2:])
        elif re.match(r'^\d+\. ', line):
            # Check if this is a TOC main item (like "1. [Overview](#overview)")
            toc_match = re.match(r'^(\d+)\.\s+\[([^\]]+)\]', line)
            if toc_match:
                number = toc_match.group(1)
                title = toc_match.group(2)
                add_toc_main_item(doc, f"{number}. {title}")
            else:
                # Regular numbered list
                match = re.match(r'^\d+\. (.+)', line)
                if match:
                    add_numbered_point(doc, match.group(1))
        
        # Handle horizontal rules
        elif line.strip() == '---':
            add_horizontal_rule(doc)
        
        # Handle regular paragraphs
        elif line:
            add_formatted_paragraph(doc, line)
        
        # Handle empty lines
        else:
            if i > 0 and lines[i-1].strip():  # Only add space after content
                doc.add_paragraph('')
        
        i += 1
    
    # Handle any remaining table at end of document
    if in_table and table_headers and 'table_data' in locals():
        add_formatted_table(doc, table_headers, table_data)
    
    doc.save(docx_path)

def add_styled_heading(doc, text, level):
    """Add a professionally styled heading"""
    heading = doc.add_heading('', level=level)
    run = heading.runs[0] if heading.runs else heading.add_run()
    run.text = text
    
    # Apply custom styling based on level
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
        run.bold = True
    elif level == 2:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0, 51, 102)
        run.bold = True
    elif level == 3:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0, 51, 102)
        run.bold = True
    
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

def add_formatted_table(doc, headers, data):
    """Add a professionally formatted table"""
    if not headers or not data:
        return
    
    # Create table
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    
    # Add headers with simple bold formatting
    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = header
        # Style header cells - just bold, no background color
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(11)
    
    # Add data rows
    for row_data in data:
        row = table.add_row()
        for i, cell_data in enumerate(row_data):
            if i < len(row.cells):
                cell = row.cells[i]
                # Handle <br> tags in cell content
                cell_text = str(cell_data)
                
                # Clear the default paragraph and add formatted content
                cell.paragraphs[0].clear()
                
                # Split by <br> tags and add as separate lines
                lines = re.split(r'<br\s*/?>', cell_text)
                
                for line_idx, line in enumerate(lines):
                    if line_idx > 0:
                        # Add a new paragraph for subsequent lines
                        new_para = cell.add_paragraph()
                    else:
                        # Use the existing paragraph for the first line
                        new_para = cell.paragraphs[0]
                    
                    # Add the line content with formatting
                    add_formatted_text_to_cell_paragraph(new_para, line.strip())
    
    # Add space after table
    doc.add_paragraph('')

def add_formatted_text_to_cell_paragraph(paragraph, text):
    """Add formatted text to a table cell paragraph"""
    if not text:
        return
    
    # Handle bold and italic formatting within the text
    parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)', text)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.size = Pt(10)
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
            run.font.size = Pt(10)
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
        else:
            if part:
                run = paragraph.add_run(part)
                run.font.size = Pt(10)

def add_formatted_paragraph(doc, text):
    """Add a formatted paragraph with text styling"""
    paragraph = doc.add_paragraph()
    
    # Handle <br> tags first by splitting text into lines
    lines = re.split(r'<br\s*/?>', text)
    
    for line_idx, line in enumerate(lines):
        if line_idx > 0:
            # Add a line break for subsequent lines
            paragraph.add_run().add_break()
        
        # Handle bold and italic formatting within each line
        parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)', line)
        
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                # Bold text
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('*') and part.endswith('*'):
                # Italic text
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            elif part.startswith('`') and part.endswith('`'):
                # Inline code
                run = paragraph.add_run(part[1:-1])
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(199, 37, 78)  # Red color for code
            else:
                # Regular text
                if part:
                    paragraph.add_run(part)
    
    # Set paragraph formatting
    paragraph.space_after = Pt(6)

def add_bullet_point(doc, text):
    """Add a bullet point"""
    paragraph = doc.add_paragraph(style='List Bullet')
    add_formatted_text_to_paragraph(paragraph, text)

def add_numbered_point(doc, text):
    """Add a numbered list item"""
    paragraph = doc.add_paragraph(style='List Number')
    add_formatted_text_to_paragraph(paragraph, text)

def add_toc_main_item(doc, text):
    """Add a main table of contents item (like '1. Overview')"""
    paragraph = doc.add_paragraph()
    # Parse the main TOC item
    toc_match = re.match(r'^(\d+\.)\s+(.+)', text)
    if toc_match:
        number = toc_match.group(1)
        title = toc_match.group(2)
        
        # Add the number part
        number_run = paragraph.add_run(number + ' ')
        number_run.font.bold = True
        number_run.font.size = Pt(12)
        number_run.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
        
        # Add the title part
        title_run = paragraph.add_run(title)
        title_run.font.size = Pt(12)
        title_run.font.bold = True
        
        # Set paragraph formatting
        paragraph.paragraph_format.left_indent = Inches(0.0)
        paragraph.space_after = Pt(6)
        paragraph.space_before = Pt(3)
    else:
        # Fallback to regular formatting
        run = paragraph.add_run(text)
        run.font.size = Pt(12)
        run.font.bold = True

def add_toc_item(doc, text):
    """Add a table of contents item with proper formatting"""
    paragraph = doc.add_paragraph()
    # Parse the TOC item (like "1.1 Introduction(#introduction)")
    toc_match = re.match(r'^(\d+\.?\d*)\s+(.+?)(\(#.+\))?$', text)
    if toc_match:
        number = toc_match.group(1)
        title = toc_match.group(2)
        
        # Add the number part
        number_run = paragraph.add_run(number + ' ')
        number_run.font.bold = True
        number_run.font.size = Pt(11)
        
        # Add the title part
        title_run = paragraph.add_run(title)
        title_run.font.size = Pt(11)
        
        # Set paragraph formatting
        paragraph.paragraph_format.left_indent = Inches(0.3)
        paragraph.space_after = Pt(3)
    else:
        # Fallback to regular formatting
        run = paragraph.add_run(text)
        run.font.size = Pt(11)

def add_toc_sub_item(doc, text):
    """Add a table of contents sub-item with indentation"""
    paragraph = doc.add_paragraph()
    # Parse the TOC sub-item
    toc_match = re.match(r'^(\d+\.\d+)\s+(.+?)(\(#.+\))?$', text)
    if toc_match:
        number = toc_match.group(1)
        title = toc_match.group(2)
        
        # Add the number part
        number_run = paragraph.add_run(number + ' ')
        number_run.font.bold = True
        number_run.font.size = Pt(10)
        number_run.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
        
        # Add the title part
        title_run = paragraph.add_run(title)
        title_run.font.size = Pt(10)
        
        # Set paragraph formatting with more indentation for sub-items
        paragraph.paragraph_format.left_indent = Inches(0.5)
        paragraph.space_after = Pt(2)
    else:
        # Fallback to regular formatting with indentation
        run = paragraph.add_run(text)
        run.font.size = Pt(10)
        paragraph.paragraph_format.left_indent = Inches(0.5)

def add_formatted_text_to_paragraph(paragraph, text):
    """Add formatted text to an existing paragraph"""
    # Handle <br> tags first by splitting text into lines
    lines = re.split(r'<br\s*/?>', text)
    
    for line_idx, line in enumerate(lines):
        if line_idx > 0:
            # Add a line break for subsequent lines
            paragraph.add_run().add_break()
        
        # Handle bold and italic formatting within each line
        parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)', line)
        
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('*') and part.endswith('*'):
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            elif part.startswith('`') and part.endswith('`'):
                run = paragraph.add_run(part[1:-1])
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
            else:
                if part:
                    paragraph.add_run(part)

def add_code_block(doc, code_text):
    """Add a formatted code block"""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0, 0, 0)
    
    # Set paragraph style for code block
    paragraph.style = doc.styles['Normal']
    paragraph.space_before = Pt(6)
    paragraph.space_after = Pt(6)
    
    # Add background color (light gray)
    paragraph_format = paragraph.paragraph_format
    paragraph_format.left_indent = Inches(0.5)
    paragraph_format.right_indent = Inches(0.5)

def add_horizontal_rule(doc):
    """Add a horizontal rule"""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.add_break(WD_BREAK.LINE)

def set_cell_background_color(cell, color):
    """Set background color for a table cell"""
    cell_xml_element = cell._tc
    table_cell_properties = cell_xml_element.get_or_add_tcPr()
    shade_obj = OxmlElement('w:shd')
    # Convert RGB to hex format
    hex_color = f"{color.r:02x}{color.g:02x}{color.b:02x}"
    shade_obj.set(qn('w:fill'), hex_color)
    table_cell_properties.append(shade_obj)

def read_docx(docx_path):
    doc = Document(docx_path)
    return '\n'.join([p.text for p in doc.paragraphs])

if __name__ == "__main__":
    # Convert the summary markdown to docx
    with open(r"C:\team of agents\agents_factory\factory_agent_summary.md", "r", encoding="utf-8") as f:
        md_text = f.read()
    create_docx_from_md(md_text, r"C:\team of agents\agents_factory\factory_agent_summary.docx")
