"""
Strategy Generator Utilities for Validation Strategist Agent
Generates structured validation strategies and reports using DOCX templates
"""

import os
import yaml
from datetime import datetime
from .document_processor import get_feature_filename, ensure_output_directory, ensure_document_libraries
from .docx_utils import write_docx, create_validation_strategy_docx

def generate_validation_strategy_from_template(feature_name, strategy_data, output_folder):
    """
    Generate validation strategy documents using both YAML and DOCX templates.
    
    Args:
        feature_name (str): Name of the feature
        strategy_data (dict): Strategy data to populate templates
        output_folder (str): Output folder path
    
    Returns:
        dict: Paths to generated documents
    """
    # Ensure libraries are available
    libraries = ensure_document_libraries()
    if not libraries.get('python-docx', False):
        print("⚠️  DOCX generation may be limited without python-docx")
    
    # Ensure output directory exists
    ensure_output_directory(output_folder)
    
    # Generate filenames
    md_filename = get_feature_filename('validation_strategy', feature_name, 'md')
    docx_filename = get_feature_filename('validation_strategy', feature_name, 'docx')
    
    md_path = os.path.join(output_folder, md_filename)
    docx_path = os.path.join(output_folder, docx_filename)
    
    results = {
        'md_path': md_path,
        'docx_path': docx_path,
        'md_success': False,
        'docx_success': False
    }
    
    # Generate Markdown version
    try:
        md_content = generate_validation_strategy_md(feature_name, strategy_data)
        with open(md_path, 'w', encoding='utf-8') as f:
            f.write(md_content)
        results['md_success'] = True
        print(f"✅ Markdown validation strategy generated: {md_path}")
    except Exception as e:
        print(f"❌ Error generating Markdown strategy: {str(e)}")
    
    # Generate DOCX version using template
    try:
        if libraries.get('python-docx', False):
            results['docx_success'] = generate_validation_strategy_docx_from_template(
                feature_name, strategy_data, docx_path
            )
        else:
            print("⚠️  Skipping DOCX generation - python-docx not available")
    except Exception as e:
        print(f"❌ Error generating DOCX strategy: {str(e)}")
    
    return results

def generate_validation_strategy_docx_from_template(feature_name, strategy_data, output_path):
    """
    Generate DOCX validation strategy using the professional template.
    
    Args:
        feature_name (str): Name of the feature
        strategy_data (dict): Strategy data
        output_path (str): Output file path
    
    Returns:
        bool: True if successful, False otherwise
    """
    try:
        from docx import Document
        
        # Load the DOCX template
        template_path = os.path.join(
            os.path.dirname(__file__), 
            '..', 'templates', 
            'validation-strategy-tmpl.docx'
        )
        
        if not os.path.exists(template_path):
            print(f"❌ Template not found: {template_path}")
            return False
        
        # Open template document
        doc = Document(template_path)
        
        current_date = datetime.now().strftime('%Y-%m-%d')
        
        # Replace placeholders in the document
        replacements = {
            '[FEATURE_NAME]': feature_name,
            '[DATE]': current_date,
            '[Domain]': strategy_data.get('feature_domain', 'TBD'),
            '[Scope]': strategy_data.get('feature_scope', 'TBD'),
            '[Criticality]': strategy_data.get('feature_criticality', 'TBD'),
        }
        
        # Replace text in paragraphs
        for paragraph in doc.paragraphs:
            for placeholder, replacement in replacements.items():
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, replacement)
        
        # Replace text in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for placeholder, replacement in replacements.items():
                        if placeholder in cell.text:
                            cell.text = cell.text.replace(placeholder, replacement)
        
        # Save the document
        doc.save(output_path)
        print(f"✅ DOCX validation strategy generated: {output_path}")
        return True
        
    except Exception as e:
        print(f"❌ Error generating DOCX from template: {str(e)}")
        return False

def generate_validation_strategy_md(feature_name, strategy_data):
    """
    Generate Markdown validation strategy content.
    
    Args:
        feature_name (str): Name of the feature
        strategy_data (dict): Strategy data
    
    Returns:
        str: Markdown content
    """
    current_date = datetime.now().strftime('%Y-%m-%d')
    
    md_content = f"""# Validation Strategy for {feature_name}

**Generated:** {current_date}  
**Domain:** {strategy_data.get('feature_domain', 'TBD')}  
**Scope:** {strategy_data.get('feature_scope', 'TBD')}  
**Criticality:** {strategy_data.get('feature_criticality', 'TBD')}

## Feature Overview

{strategy_data.get('feature_description', 'Feature description to be added.')}

## Validation Approach

### Test Strategy
- **Verification Level:** {strategy_data.get('verification_level', 'Unit/Integration/System')}
- **Test Environment:** {strategy_data.get('test_environment', 'TBD')}
- **Test Automation:** {strategy_data.get('automation_level', 'TBD')}

### Coverage Areas
{format_coverage_areas(strategy_data.get('coverage_areas', []))}

### Risk Assessment
{format_risk_assessment(strategy_data.get('risks', []))}

### Test Planning
{format_test_planning(strategy_data.get('test_plan', {}))}

## Dependencies

{format_dependencies(strategy_data.get('dependencies', []))}

## Timeline & Milestones

{format_timeline(strategy_data.get('timeline', {}))}

## Success Criteria

{format_success_criteria(strategy_data.get('success_criteria', []))}

---
*Generated by Validation Strategist Agent*
"""
    return md_content

def format_coverage_areas(coverage_areas):
    """Format coverage areas for Markdown output."""
    if not coverage_areas:
        return "- Coverage areas to be defined"
    
    return "\n".join([f"- {area}" for area in coverage_areas])

def format_risk_assessment(risks):
    """Format risk assessment for Markdown output."""
    if not risks:
        return "Risk assessment to be completed."
    
    risk_text = "| Risk | Impact | Probability | Mitigation |\n"
    risk_text += "|------|--------|-------------|------------|\n"
    
    for risk in risks:
        risk_text += f"| {risk.get('description', 'TBD')} | {risk.get('impact', 'TBD')} | {risk.get('probability', 'TBD')} | {risk.get('mitigation', 'TBD')} |\n"
    
    return risk_text

def format_test_planning(test_plan):
    """Format test planning for Markdown output."""
    if not test_plan:
        return "Test planning details to be added."
    
    planning_text = ""
    if test_plan.get('phases'):
        planning_text += "### Test Phases\n"
        for phase in test_plan['phases']:
            planning_text += f"- **{phase.get('name', 'TBD')}:** {phase.get('description', 'TBD')}\n"
    
    if test_plan.get('resources'):
        planning_text += "\n### Resources Required\n"
        planning_text += f"- {test_plan['resources']}\n"
    
    return planning_text if planning_text else "Test planning details to be added."

def format_dependencies(dependencies):
    """Format dependencies for Markdown output."""
    if not dependencies:
        return "No external dependencies identified."
    
    return "\n".join([f"- {dep}" for dep in dependencies])

def format_timeline(timeline):
    """Format timeline for Markdown output."""
    if not timeline:
        return "Timeline to be established."
    
    timeline_text = "| Milestone | Target Date | Status |\n"
    timeline_text += "|-----------|-------------|--------|\n"
    
    for milestone, details in timeline.items():
        if isinstance(details, dict):
            target = details.get('target_date', 'TBD')
            status = details.get('status', 'Planned')
        else:
            target = str(details)
            status = 'Planned'
        
        timeline_text += f"| {milestone} | {target} | {status} |\n"
    
    return timeline_text

def format_success_criteria(success_criteria):
    """Format success criteria for Markdown output."""
    if not success_criteria:
        return "Success criteria to be defined."
    
    return "\n".join([f"- {criteria}" for criteria in success_criteria])

class ValidationStrategyGenerator:
    """Main class for generating validation strategies and related documents."""
    
    def __init__(self, agent_root_path):
        """
        Initialize the strategy generator.
        
        Args:
            agent_root_path (str): Root path of the validation strategist agent
        """
        self.agent_root = agent_root_path
        self.output_path = os.path.join(agent_root_path, "validation strategy")
        self.templates_path = os.path.join(agent_root_path, ".bmad-core", "templates")
        
    def generate_validation_strategy(self, feature_name, strategy_data):
        """
        Generate a complete validation strategy document.
        
        Args:
            feature_name (str): Name of the feature
            strategy_data (dict): Strategy data structure
        
        Returns:
            dict: Generated file paths and status
        """
        result = {
            "success": False,
            "md_file": None,
            "docx_file": None,
            "errors": []
        }
        
        try:
            # Ensure output directory exists
            ensure_output_directory(self.output_path)
            
            # Generate filenames
            md_filename = get_feature_filename("validation_strategy", feature_name, "md")
            docx_filename = get_feature_filename("validation_strategy", feature_name, "docx")
            
            md_file_path = os.path.join(self.output_path, md_filename)
            docx_file_path = os.path.join(self.output_path, docx_filename)
            
            # Generate Markdown content
            md_content = self._create_strategy_markdown(feature_name, strategy_data)
            
            # Write Markdown file
            with open(md_file_path, 'w', encoding='utf-8') as f:
                f.write(md_content)
            result["md_file"] = md_file_path
            print(f"✅ Validation strategy MD created: {md_filename}")
            
            # Write DOCX file
            if create_validation_strategy_docx(strategy_data, docx_file_path, feature_name):
                result["docx_file"] = docx_file_path
            else:
                result["errors"].append("Failed to create DOCX file")
            
            result["success"] = True
            return result
            
        except Exception as e:
            result["errors"].append(f"Strategy generation error: {str(e)}")
            return result
    
    def generate_risk_assessment(self, feature_name, risk_data):
        """
        Generate a risk assessment document.
        
        Args:
            feature_name (str): Name of the feature
            risk_data (dict): Risk assessment data
        
        Returns:
            dict: Generated file paths and status
        """
        result = {
            "success": False,
            "md_file": None,
            "docx_file": None,
            "errors": []
        }
        
        try:
            ensure_output_directory(self.output_path)
            
            # Generate filenames
            md_filename = get_feature_filename("risk_assessment", feature_name, "md")
            docx_filename = get_feature_filename("risk_assessment", feature_name, "docx")
            
            md_file_path = os.path.join(self.output_path, md_filename)
            docx_file_path = os.path.join(self.output_path, docx_filename)
            
            # Generate content
            md_content = self._create_risk_assessment_markdown(feature_name, risk_data)
            
            # Write files
            with open(md_file_path, 'w', encoding='utf-8') as f:
                f.write(md_content)
            result["md_file"] = md_file_path
            
            if write_docx(md_content, docx_file_path, f"Risk Assessment: {feature_name}"):
                result["docx_file"] = docx_file_path
            
            result["success"] = True
            print(f"✅ Risk assessment created: {md_filename}")
            return result
            
        except Exception as e:
            result["errors"].append(f"Risk assessment generation error: {str(e)}")
            return result
    
    def generate_architecture_analysis(self, feature_name, analysis_data):
        """
        Generate an architecture analysis document.
        
        Args:
            feature_name (str): Name of the feature
            analysis_data (dict): Architecture analysis data
        
        Returns:
            dict: Generated file paths and status
        """
        result = {
            "success": False,
            "md_file": None,
            "docx_file": None,
            "errors": []
        }
        
        try:
            ensure_output_directory(self.output_path)
            
            # Generate filenames
            md_filename = get_feature_filename("architecture_analysis", feature_name, "md")
            docx_filename = get_feature_filename("architecture_analysis", feature_name, "docx")
            
            md_file_path = os.path.join(self.output_path, md_filename)
            docx_file_path = os.path.join(self.output_path, docx_filename)
            
            # Generate content
            md_content = self._create_architecture_analysis_markdown(feature_name, analysis_data)
            
            # Write files
            with open(md_file_path, 'w', encoding='utf-8') as f:
                f.write(md_content)
            result["md_file"] = md_file_path
            
            if write_docx(md_content, docx_file_path, f"Architecture Analysis: {feature_name}"):
                result["docx_file"] = docx_file_path
            
            result["success"] = True
            print(f"✅ Architecture analysis created: {md_filename}")
            return result
            
        except Exception as e:
            result["errors"].append(f"Architecture analysis generation error: {str(e)}")
            return result
    
    def _create_strategy_markdown(self, feature_name, strategy_data):
        """Create Markdown content for validation strategy."""
        timestamp = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        
        content = f"""# Validation Strategy: {feature_name}

**Generated:** {timestamp}  
**Agent:** Validation Strategist (Val)  
**Document Type:** Strategic Validation Plan

## Executive Summary

{strategy_data.get('executive_summary', 'Strategic validation approach for ' + feature_name)}

## Feature Overview

{strategy_data.get('feature_overview', 'Feature overview to be provided')}

## Architecture Analysis

{strategy_data.get('architecture_analysis', 'Architecture analysis to be completed')}

### Key Architectural Elements
{strategy_data.get('architectural_elements', '- To be identified')}

### Critical Interfaces
{strategy_data.get('critical_interfaces', '- To be identified')}

### Dependencies
{strategy_data.get('dependencies', '- To be analyzed')}

## Validation Strategy

### Strategic Approach
{strategy_data.get('strategic_approach', 'Strategic validation approach to be defined')}

### Coverage Strategy
{strategy_data.get('coverage_strategy', 'Coverage strategy to be developed')}

### Priority Areas
{strategy_data.get('priority_areas', '- To be prioritized based on risk assessment')}

## Risk Assessment

### High-Level Risks
{strategy_data.get('high_level_risks', '- To be identified and assessed')}

### Mitigation Strategies
{strategy_data.get('mitigation_strategies', '- To be developed')}

## Resource Strategy

### Required Skills
{strategy_data.get('required_skills', '- To be determined')}

### Tool Requirements
{strategy_data.get('tool_requirements', '- To be specified')}

### Environment Needs
{strategy_data.get('environment_needs', '- To be defined')}

### Timeline Considerations
{strategy_data.get('timeline_considerations', '- To be estimated')}

## Success Criteria

{strategy_data.get('success_criteria', 'Success criteria to be defined')}

## Next Steps

{strategy_data.get('next_steps', '1. Complete detailed analysis\\n2. Develop tactical plans\\n3. Begin implementation')}

---
*This document was generated by the Validation Strategist Agent as part of the strategic validation planning process.*
"""
        return content
    
    def _create_risk_assessment_markdown(self, feature_name, risk_data):
        """Create Markdown content for risk assessment."""
        timestamp = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        
        content = f"""# Risk Assessment: {feature_name}

**Generated:** {timestamp}  
**Agent:** Validation Strategist (Val)  
**Document Type:** Strategic Risk Assessment

## Executive Summary

{risk_data.get('executive_summary', 'Strategic risk assessment for ' + feature_name)}

## Risk Identification

### Technical Risks
{risk_data.get('technical_risks', '- To be identified')}

### Integration Risks
{risk_data.get('integration_risks', '- To be analyzed')}

### Schedule Risks
{risk_data.get('schedule_risks', '- To be assessed')}

### Resource Risks
{risk_data.get('resource_risks', '- To be evaluated')}

## Risk Analysis

### Risk Priority Matrix
{risk_data.get('risk_matrix', '| Risk | Probability | Impact | Priority |\\n|------|-------------|--------|----------|\\n| TBD  | TBD         | TBD    | TBD      |')}

### Critical Path Risks
{risk_data.get('critical_path_risks', '- To be identified')}

## Mitigation Strategies

### Prevention Strategies
{risk_data.get('prevention_strategies', '- To be developed')}

### Contingency Plans
{risk_data.get('contingency_plans', '- To be created')}

### Monitoring Approach
{risk_data.get('monitoring_approach', '- To be established')}

## Risk Monitoring Plan

{risk_data.get('monitoring_plan', 'Risk monitoring plan to be developed')}

## Recommendations

{risk_data.get('recommendations', 'Strategic recommendations to be provided')}

---
*This risk assessment was generated by the Validation Strategist Agent.*
"""
        return content
    
    def _create_architecture_analysis_markdown(self, feature_name, analysis_data):
        """Create Markdown content for architecture analysis."""
        timestamp = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        
        content = f"""# Architecture Analysis: {feature_name}

**Generated:** {timestamp}  
**Agent:** Validation Strategist (Val)  
**Document Type:** Strategic Architecture Analysis

## Executive Summary

{analysis_data.get('executive_summary', 'Architecture analysis for ' + feature_name)}

## System Architecture Overview

{analysis_data.get('system_overview', 'System architecture overview to be provided')}

## Component Analysis

### Key Components
{analysis_data.get('key_components', '- To be identified')}

### Component Interactions
{analysis_data.get('component_interactions', '- To be mapped')}

## Interface Analysis

### External Interfaces
{analysis_data.get('external_interfaces', '- To be documented')}

### Internal Interfaces
{analysis_data.get('internal_interfaces', '- To be analyzed')}

## Dependency Mapping

### Feature Dependencies
{analysis_data.get('feature_dependencies', '- To be mapped')}

### System Dependencies
{analysis_data.get('system_dependencies', '- To be identified')}

## Integration Points

### Critical Integration Points
{analysis_data.get('critical_integration_points', '- To be highlighted')}

### Integration Complexity
{analysis_data.get('integration_complexity', '- To be assessed')}

## Validation Implications

### Architecture-Driven Validation Requirements
{analysis_data.get('validation_requirements', '- To be derived')}

### Strategic Validation Considerations
{analysis_data.get('strategic_considerations', '- To be analyzed')}

## Recommendations

{analysis_data.get('recommendations', 'Strategic recommendations based on architecture analysis')}

---
*This architecture analysis was generated by the Validation Strategist Agent.*
"""
        return content

if __name__ == "__main__":
    # Test strategy generator
    print("🧪 Testing Strategy Generator...")
    
    # Test data
    test_strategy_data = {
        "executive_summary": "Test validation strategy for feature X",
        "feature_overview": "Feature X provides critical functionality",
        "strategic_approach": "Risk-based validation approach",
        "success_criteria": "All critical paths validated successfully"
    }
    
    # This would normally use the actual agent root path
    test_root = os.getcwd()
    generator = ValidationStrategyGenerator(test_root)
    
    result = generator.generate_validation_strategy("TestFeature", test_strategy_data)
    
    if result["success"]:
        print("✅ Strategy generator test successful!")
        print(f"   MD file: {result['md_file']}")
        print(f"   DOCX file: {result['docx_file']}")
    else:
        print("❌ Strategy generator test failed!")
        for error in result["errors"]:
            print(f"   Error: {error}")